import json
import csv
import re
from pathlib import Path
from typing import Dict, Any
import pypdf
import docx
from src.utils.ocr import extract_text_from_image

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".csv", ".json",
    ".log", ".txt",
    ".png", ".jpg", ".jpeg", ".webp"
}


def extract_file(file_path: Path, original_name: str) -> Dict[str, Any]:
    ext = Path(original_name).suffix.lower()
    metadata = {
        "original_name": original_name,
        "file_extension": ext,
        "size_bytes": file_path.stat().st_size if file_path.exists() else 0,
    }

    # Strict file type whitelist
    if ext not in ALLOWED_EXTENSIONS:
        return {
            "file_type": ext.lstrip(".") if ext else "unknown",
            "extractor": "blocked",
            "success": False,
            "error": f"unsupported file type: {ext}",
            "text": "",
            "metadata": {
                **metadata,
                "blocked": True,
                "reason": "unsupported_file_type",
            },
            "chunks": [],
        }

    text = ""
    error = None
    success = True
    extractor = "unknown"
    page_count = 0
    line_count = 0
    table_count = 0
    heading_count = 0
    chunks = []

    try:
        if ext == ".pdf":
            extractor = "pypdf"
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                page_count = len(reader.pages)
                pages_text = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    pages_text.append(page_text)
                text = "\n\n--- page break ---\n\n".join(pages_text)

        elif ext == ".docx":
            extractor = "python-docx"
            doc = docx.Document(file_path)
            lines = []

            for para in doc.paragraphs:
                style_name = (para.style.name or "").lower()
                para_text = para.text.strip()

                if not para_text:
                    continue

                if "heading" in style_name:
                    level = style_name.replace("heading", "").strip()
                    if level.isdigit():
                        lines.append(f"{'#' * int(level)} {para_text}")
                    else:
                        lines.append(f"## {para_text}")
                    heading_count += 1
                elif "list" in style_name or para_text.startswith(("•", "-", "*")):
                    lines.append(f"- {para_text}")
                else:
                    lines.append(para_text)

            for table in doc.tables:
                table_count += 1
                lines.append(f"\n**Table {table_count}**\n")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    lines.append(row_text)
                lines.append("")

            text = "\n".join(lines)
            line_count = len(lines)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = table_count
            metadata["heading_count"] = heading_count

        elif ext == ".csv":
            extractor = "csv"
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                rows = list(reader)
                line_count = len(rows)

                if rows:
                    header = rows[0]
                    text = "| " + " | ".join(header) + " |\n"
                    text += "|" + "|".join(["---"] * len(header)) + "|\n"
                    for row in rows[1:]:
                        text += "| " + " | ".join(row) + " |\n"

            metadata["row_count"] = line_count

        elif ext == ".json":
            extractor = "json"
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)
                text = json.dumps(data, indent=2, ensure_ascii=False)

            if isinstance(data, dict):
                metadata["json_keys"] = list(data.keys())
            else:
                metadata["json_keys"] = None

        elif ext in [".log", ".txt"]:
            extractor = "text"
            lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
            line_count = len(lines)
            text = "\n".join(lines)
            metadata["line_count"] = line_count

        elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
            extractor = "tesseract"
            ocr_result = extract_text_from_image(file_path)

            if ocr_result["success"]:
                text = ocr_result["text"]
                metadata["ocr_used"] = True
                metadata["ocr_text_length"] = ocr_result["metadata"]["text_length"]
                metadata["ocr_line_count"] = ocr_result["metadata"]["line_count"]
                metadata["ocr_confidence_hint"] = ocr_result["metadata"]["confidence_hint"]
                metadata["ocr_text_found"] = bool(text.strip())
            else:
                return {
                    "file_type": ext.lstrip(".") if ext else "unknown",
                    "extractor": extractor,
                    "success": False,
                    "error": ocr_result["error"],
                    "text": "",
                    "metadata": {
                        **metadata,
                        "ocr_used": False,
                        "ocr_error": ocr_result["error"],
                        "extraction_failed": True,
                    },
                    "chunks": [],
                }

        # Normalize whitespace but preserve line breaks
        text = re.sub(r"[ \t]+", " ", text)

        if len(text) > 5000:
            chunk_size = 2000
            chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
        else:
            chunks = [text] if text else []

        if ext == ".pdf":
            metadata["page_count"] = page_count

        return {
            "file_type": ext.lstrip(".") if ext else "unknown",
            "extractor": extractor,
            "success": success,
            "error": error,
            "text": text,
            "metadata": metadata,
            "chunks": chunks,
        }

    except Exception as e:
        return {
            "file_type": ext.lstrip(".") if ext else "unknown",
            "extractor": extractor,
            "success": False,
            "error": str(e),
            "text": "",
            "metadata": {
                **metadata,
                "extraction_failed": True,
            },
            "chunks": [],
        }